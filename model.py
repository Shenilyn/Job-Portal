# type: ignore
from tensorflow.keras.models import load_model  # type: ignore

import os
import numpy as np
import tensorflow as tf
import json
import requests
from dotenv import load_dotenv
import pandas as pd
import pdfplumber
import docx
import pickle

# Load environment variables (for API keys)
load_dotenv()

class JobRecommendationSystem:
    def __init__(self, model_path="model/model3.h5", vectorizer_path="vectorizer.pkl", 
                 job_mapping_path="job_titles.json"):
        """
        Initialize the Job Recommendation System.
        
        Args:
            model_path (str): Path to the trained model file
            vectorizer_path (str): Path to the trained TF-IDF vectorizer
            job_mapping_path (str): Path to the job title mapping JSON file
        """
        # Check if model file exists
        if not os.path.exists(model_path):
            print(f"Model file not found at: {model_path}")
            print("Please check if the file exists and the path is correct.")
            raise FileNotFoundError(f"Model file not found: {model_path}")
            
        # Verify model file size
        model_size = os.path.getsize(model_path)
        if model_size < 1000:  # Very small file, likely not a valid model
            print(f"Warning: Model file is very small ({model_size} bytes). It may be corrupted.")
            
        # Load the model
        try:
            print(f"Attempting to load model from {model_path}...")
            self.model = load_model(model_path)
            print(f"Model loaded successfully from {model_path}")
        except Exception as e:
            print(f"Error loading model: {e}")
            print("\nPossible issues:")
            print("1. The file is not a valid TensorFlow model")
            print("2. The model was saved with a different TensorFlow version")
            print("3. The file is corrupted")
            print("\nTry recreating your model and saving it again.")
            raise

        # Load the trained TF-IDF vectorizer
        try:
            with open(vectorizer_path, "rb") as f:
                self.vectorizer = pickle.load(f)
            print("TF-IDF vectorizer loaded successfully.")
        except Exception as e:
            print(f"Error loading vectorizer: {e}")
            raise
            
        # Load job title mapping
        try:
            if os.path.exists(job_mapping_path):
                with open(job_mapping_path, "r") as f:
                    self.job_titles = json.load(f)
                print(f"Job titles mapping loaded successfully from {job_mapping_path}")
            else:
                # Create a default mapping if file doesn't exist
                print(f"Job mapping file not found at {job_mapping_path}. Using default indices.")
                self.job_titles = {str(i): f"Job Category {i+1}" for i in range(self.model.output_shape[1])}
        except Exception as e:
            print(f"Error loading job titles: {e}")
            # Fallback mapping
            self.job_titles = {str(i): f"Job Category {i+1}" for i in range(self.model.output_shape[1])}

        # Initialize API keys and endpoints
        self.gemini_api_key = os.environ.get("VITE_GEMINI_API_KEY")
        if not self.gemini_api_key:
            print("Warning: VITE_GEMINI_API_KEY not found in environment variables")
        self.gemini_api_url = "https://generativelanguage.googleapis.com/v1beta/models/gemini-pro:generateContent"

    def extract_text_from_resume(self, file_path):
        """
        Extract text from PDF or DOCX resume files.
        
        Args:
            file_path (str): Path to the resume file
            
        Returns:
            str: Extracted text from the resume
            
        Raises:
            ValueError: If the file format is not supported
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Resume file not found: {file_path}")
            
        if file_path.lower().endswith(".pdf"):
            try:
                with pdfplumber.open(file_path) as pdf:
                    return "\n".join(page.extract_text() or "" for page in pdf.pages)
            except Exception as e:
                raise ValueError(f"Error extracting text from PDF: {e}")
        elif file_path.lower().endswith(".docx"):
            try:
                doc = docx.Document(file_path)
                return "\n".join(paragraph.text for paragraph in doc.paragraphs)
            except Exception as e:
                raise ValueError(f"Error extracting text from DOCX: {e}")
        else:
            raise ValueError("Unsupported file format. Only .pdf and .docx are supported.")

    def preprocess_resume(self, resume_text):
        """
        Preprocess resume text using the trained TF-IDF vectorizer.
        
        Args:
            resume_text (str): Raw text extracted from the resume
            
        Returns:
            numpy.ndarray: Feature vector for model prediction
        """
        try:
            features = self.vectorizer.transform([resume_text]).toarray()
            return features
        except Exception as e:
            print(f"Error during preprocessing: {e}")
            # Return zeros array matching the expected input shape
            return np.zeros((1, self.model.input_shape[1]))

    def get_top_job_recommendations(self, resume_text, top_n=3):
        """
        Get top job recommendations based on resume text.
        
        Args:
            resume_text (str): Preprocessed resume text
            top_n (int): Number of top recommendations to return
            
        Returns:
            list: List of dictionaries containing job titles and confidence scores
        """
        features = self.preprocess_resume(resume_text)
        
        try:
            predictions = self.model.predict(features)
            # Ensure top_n doesn't exceed the number of job categories
            top_n = min(top_n, len(self.job_titles))
            # Get indices of top predictions
            top_indices = np.argsort(predictions[0])[-top_n:][::-1]
            
            return [
                {
                    "job_title": self.job_titles.get(str(idx), f"Job Category {idx+1}"),
                    "confidence": float(predictions[0][idx])
                }
                for idx in top_indices
            ]
        except Exception as e:
            print(f"Error during prediction: {e}")
            return [{"job_title": "Error in prediction", "confidence": 0.0}]

    def get_training_courses(self, job_title, resume_text):
        """
        Get recommended training courses for a specific job title using Gemini API.
        
        Args:
            job_title (str): The job title to get courses for
            resume_text (str): Resume text to provide context
            
        Returns:
            list: List of recommended courses
        """
        if not self.gemini_api_key:
            return [{"course_name": "API Key Missing", 
                    "provider": "N/A", 
                    "description": "No Gemini API key provided in environment variables.",
                    "url": "",
                    "relevance": "Please set the VITE_GEMINI_API_KEY environment variable."}]
        
        prompt = f"""
        Based on this resume: 
        
        {resume_text}
        
        I need 3 specific training courses available online that would help this person qualify for a {job_title} position.
        
        For each course, provide:
        1. Course name
        2. Provider (website/platform)
        3. Brief description of what skills it will teach
        4. URL if available
        5. Why it's relevant for this specific job
        
        Format as a JSON list with course name, provider, description, url, and relevance fields.
        """
        
        headers = {
            "Content-Type": "application/json",
            "x-goog-api-key": self.gemini_api_key
        }
        
        data = {
            "contents": [{"parts": [{"text": prompt}]}]
        }
        
        try:
            response = requests.post(self.gemini_api_url, headers=headers, json=data)
            if response.status_code == 200:
                response_data = response.json()
                text_response = response_data["candidates"][0]["content"]["parts"][0]["text"]
                try:
                    # Try to extract JSON from the response
                    json_start = text_response.find('[')
                    json_end = text_response.rfind(']') + 1
                    if json_start >= 0 and json_end > 0:
                        json_content = text_response[json_start:json_end]
                        return json.loads(json_content)
                    else:
                        return self._extract_courses_from_text(text_response)
                except json.JSONDecodeError:
                    return self._extract_courses_from_text(text_response)
            else:
                error_msg = f"Error calling Gemini API: {response.status_code}"
                try:
                    error_detail = response.json().get("error", {}).get("message", "Unknown error")
                    error_msg += f" - {error_detail}"
                except:
                    error_msg += f" - {response.text[:100]}"
                    
                return [{
                    "course_name": "API Error", 
                    "provider": "N/A", 
                    "description": error_msg,
                    "url": "",
                    "relevance": "Please check your API key and network connection."
                }]
        except Exception as e:
            return [{
                "course_name": "Exception", 
                "provider": "N/A", 
                "description": f"Exception when calling Gemini API: {str(e)}",
                "url": "",
                "relevance": "Please check your network connection."
            }]

    def _extract_courses_from_text(self, text):
        """
        Extract course information from non-JSON formatted text.
        
        Args:
            text (str): Text response from Gemini API
            
        Returns:
            list: List of course dictionaries
        """
        courses = []
        sections = []
        current_section = ""
        lines = text.split('\n')
        
        # Split text into course sections
        for line in lines:
            if line.strip().startswith(('1.', '2.', '3.', 'Course 1:', 'Course 2:', 'Course 3:')):
                if current_section:
                    sections.append(current_section)
                current_section = line
            elif current_section:
                current_section += "\n" + line
        
        if current_section:
            sections.append(current_section)
        
        # Extract course information from each section
        for section in sections:
            course = {
                "course_name": "Unknown Course",
                "provider": "Unknown Provider",
                "description": section,
                "url": "",
                "relevance": ""
            }
            
            for line in section.split('\n'):
                line = line.strip()
                if "course" in line.lower() or "name" in line.lower():
                    parts = line.split(':', 1)
                    if len(parts) > 1:
                        course["course_name"] = parts[1].strip()
                elif "provider" in line.lower():
                    parts = line.split(':', 1)
                    if len(parts) > 1:
                        course["provider"] = parts[1].strip()
                elif "description" in line.lower():
                    parts = line.split(':', 1)
                    if len(parts) > 1:
                        course["description"] = parts[1].strip()
                elif "url" in line.lower() or "http" in line.lower():
                    parts = line.split(':', 1)
                    if len(parts) > 1:
                        course["url"] = parts[1].strip()
                elif "relevan" in line.lower():
                    parts = line.split(':', 1)
                    if len(parts) > 1:
                        course["relevance"] = parts[1].strip()
            
            courses.append(course)
        
        return courses[:3]  # Return at most 3 courses

    def process_resume_file(self, file_path):
        """
        Process a resume file to get job recommendations and training courses.
        
        Args:
            file_path (str): Path to the resume file
            
        Returns:
            dict: Dictionary containing job recommendations and training courses
        """
        try:
            resume_text = self.extract_text_from_resume(file_path)
            job_recommendations = self.get_top_job_recommendations(resume_text)
            
            results = []
            for job in job_recommendations:
                job_title = job["job_title"]
                courses = self.get_training_courses(job_title, resume_text)
                results.append({
                    "job_title": job_title,
                    "confidence": job["confidence"],
                    "training_courses": courses
                })
            
            return {"recommendations": results}
        except Exception as e:
            return {"error": str(e)}


# Example usage
if __name__ == "__main__":
    # Example job titles mapping (you should create a proper mapping file)
    example_job_mapping = {
        "0": "Data Scientist",
        "1": "Software Engineer",
        "2": "Product Manager",
        "3": "UX Designer",
        "4": "DevOps Engineer"
    }
    
    # Save example mapping if it doesn't exist
    if not os.path.exists("job_titles.json"):
        with open("job_titles.json", "w") as f:
            json.dump(example_job_mapping, f, indent=2)
    
    print("\n=== Job Recommendation System ===")
    print("This script requires three files to run properly:")
    print("1. A TensorFlow model file (.h5)")
    print("2. A trained TF-IDF vectorizer (.pkl)")
    print("3. A resume file to analyze (.pdf or .docx)")
    print("\nIf you're getting errors, try the fallback mode which bypasses model loading.")
    
    use_fallback = input("\nDo you want to use fallback mode? (yes/no): ").lower().startswith('y')
    
    if use_fallback:
        print("\n=== Running in Fallback Mode ===")
        
        # Scan for PDF and DOCX files in the current directory and subdirectories
        print("\nScanning for resume files (PDF/DOCX)...")
        resume_files = []
        for root, _, files in os.walk('.'):
            for file in files:
                if file.lower().endswith(('.pdf', '.docx')):
                    file_path = os.path.join(root, file)
                    resume_files.append(file_path)
        
        if resume_files:
            print("\nFound the following resume files:")
            for i, file_path in enumerate(resume_files):
                print(f"{i+1}. {file_path}")
            
            try:
                choice = int(input("\nEnter the number of the file to analyze (or 0 to enter path manually): "))
                if 1 <= choice <= len(resume_files):
                    resume_file_path = resume_files[choice-1]
                else:
                    resume_file_path = input("Enter the path to the resume file to analyze: ")
            except ValueError:
                resume_file_path = input("Enter the path to the resume file to analyze: ")
        else:
            print("No PDF or DOCX files found in the current directory or subdirectories.")
            resume_file_path = input("Enter the path to the resume file to analyze: ")
        
        # Verify file exists
        if not os.path.exists(resume_file_path):
            print(f"\nERROR: File not found: {resume_file_path}")
            print("Please make sure the file exists and the path is correct.")
            print("Would you like to create a dummy resume file for testing? (yes/no)")
            create_dummy = input().lower().startswith('y')
            
            if create_dummy:
                dummy_resume_path = "./dummy_resume.txt"
                with open(dummy_resume_path, "w") as f:
                    f.write("""
JOHN DOE
Software Engineer
john.doe@example.com | (123) 456-7890 | linkedin.com/in/johndoe

SUMMARY
Experienced software engineer with 5+ years of experience in Python, JavaScript, and machine learning technologies.
Skilled in developing REST APIs, web applications, and data processing pipelines.

EXPERIENCE
Senior Software Engineer | TechCorp Inc. | 2023-Present
- Developed a machine learning pipeline that improved data processing efficiency by 35%
- Implemented RESTful APIs using Flask and FastAPI
- Led a team of 3 junior developers on a customer-facing web application project

Software Engineer | DataSoft Solutions | 2020-2023
- Created Python scripts for data analysis and visualization
- Built a React-based dashboard for real-time data monitoring
- Collaborated with data scientists to implement ML models in production

EDUCATION
Master of Science in Computer Science | State University | 2020
Bachelor of Science in Software Engineering | Tech Institute | 2018

SKILLS
- Programming: Python, JavaScript, Java, SQL
- Frameworks: React, Flask, FastAPI, TensorFlow
- Tools: Git, Docker, Kubernetes, AWS
- Soft Skills: Team Leadership, Project Management, Communication
                    """)
                resume_file_path = dummy_resume_path
                print(f"Created dummy resume at {dummy_resume_path}")
            else:
                print("Exiting program due to missing file.")
                exit(1)
        
        # Create a minimal system that doesn't use the actual model
        class FallbackSystem:
            def extract_text_from_resume(self, file_path):
                if file_path.lower().endswith(".pdf"):
                    try:
                        with pdfplumber.open(file_path) as pdf:
                            return "\n".join(page.extract_text() or "" for page in pdf.pages)
                    except Exception as e:
                        return f"Error extracting text: {e}"
                elif file_path.lower().endswith(".docx"):
                    try:
                        doc = docx.Document(file_path)
                        return "\n".join(paragraph.text for paragraph in doc.paragraphs)
                    except Exception as e:
                        return f"Error extracting text: {e}"
                elif file_path.lower().endswith(".txt"):
                    try:
                        with open(file_path, 'r') as f:
                            return f.read()
                    except Exception as e:
                        return f"Error reading text file: {e}"
                else:
                    return "Unsupported file format. Only .pdf, .docx, and .txt are supported."
            
            def process_resume_file(self, file_path):
                resume_text = self.extract_text_from_resume(file_path)
                
                # Basic keyword analysis for fallback job matching
                keywords = {
                    "Data Scientist": ["data science", "machine learning", "python", "statistics", "analytics", "data analysis", "pandas", "numpy", "sklearn"],
                    "Software Engineer": ["software", "programming", "development", "java", "python", "javascript", "code", "algorithm", "api"],
                    "Product Manager": ["product", "management", "strategy", "roadmap", "agile", "scrum", "user experience", "prioritization"],
                    "UX Designer": ["design", "user experience", "ux", "ui", "wireframe", "prototype", "usability", "sketch", "figma"],
                    "DevOps Engineer": ["devops", "ci/cd", "pipeline", "aws", "cloud", "docker", "kubernetes", "infrastructure"]
                }
                
                # Count keyword matches for each job
                job_scores = {}
                resume_text_lower = resume_text.lower()
                for job, terms in keywords.items():
                    score = sum(resume_text_lower.count(term.lower()) for term in terms)
                    job_scores[job] = score
                
                # Sort by score
                sorted_jobs = sorted(job_scores.items(), key=lambda x: x[1], reverse=True)
                
                # Create recommendations
                recommendations = []
                for job, score in sorted_jobs[:3]:  # Top 3 jobs
                    # Normalize score to a confidence between 0.5 and 0.95
                    confidence = min(0.95, max(0.5, 0.5 + (score / 10)))
                    
                    recommendations.append({
                        "job_title": job,
                        "confidence": round(confidence, 2),
                        "training_courses": [
                            {
                                "course_name": f"{job} Fundamentals",
                                "provider": "Coursera",
                                "description": f"Comprehensive training for {job} roles",
                                "url": "https://www.coursera.org",
                                "relevance": "Core professional skills"
                            }
                        ]
                    })
                
                return {
                    "resume_text": resume_text[:500] + ("..." if len(resume_text) > 500 else ""),
                    "recommendations": recommendations
                }
        
        try:
            system = FallbackSystem()
            result = system.process_resume_file(resume_file_path)
            print("\nExtracted text sample from resume:")
            print(result["resume_text"])
            print("\nRecommendations (based on basic keyword matching):")
            print(json.dumps({"recommendations": result["recommendations"]}, indent=2))
            print("\nNOTE: These recommendations use basic keyword matching only. The ML model was not used.")
        except Exception as e:
            print(f"Error in fallback mode: {e}")
    else:
        # Set the correct paths for your files
        model_path = input("Enter the path to your model file (e.g., ./your_model.h5): ./model3.h5")
        vectorizer_path = input("Enter the path to your vectorizer file (e.g., ./vectorizer.pkl): ./vectorizer.py")
        resume_file_path = input("Enter the path to the resume file to analyze: ./10553553.pdf")
        
        try:
            job_system = JobRecommendationSystem(
                model_path=model_path,
                vectorizer_path=vectorizer_path
            )
            result = job_system.process_resume_file(resume_file_path)
            print(json.dumps(result, indent=2))
        except Exception as e:
            print(f"Error: {e}")
            print("\nTry running the script again and selecting fallback mode to test resume parsing.")